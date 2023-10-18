import re
import tkinter as tk
from tkinter import simpledialog
from docx import Document
from docx.shared import Pt

month_to_number = {
    'Jan': '01', 'Feb': '02', 'Mar': '03', 'Apr': '04',
    'May': '05', 'Jun': '06', 'Jul': '07', 'Aug': '08',
    'Sep': '09', 'Oct': '10', 'Nov': '11', 'Dec': '12'
}

write_data = []

def add_tweet():
    tweet = tweet_entry.get()
    url = url_entry.get()  # Get the URL input
    comment = comment_entry.get()  # Get the comment input
    if tweet:
        tweets.append(tweet)  # Store the tweet text
        urls.append(url)  # Store the URL
        comments.append(comment)  # Store the comment
        tweet_entry.delete(0, tk.END)
        url_entry.delete(0, tk.END)  # Clear the URL input field
        comment_entry.delete(0, tk.END)  # Clear the comment input field
        update_tweet_list()

def delete_selected_item(event):
    selected_index = tweet_listbox.curselection()  # Get the index of the selected item(s)
    if selected_index:
        for index in selected_index:
            index = int(index)
            del tweets[index]  # Delete the selected item from the 'tweets' list
            del urls[index]  # Delete the selected item from the 'urls' list
            del comments[index]  # Delete the selected item from the 'comments' list
            tweet_listbox.delete(index)  # Delete the selected item from the Listbox
            update_tweet_list()  # Update the Listbox to reflect the changes

def update_tweet_list():
    tweet_listbox.delete(0, tk.END)
    for i, tweet in enumerate(tweets):
        tweet_listbox.insert(tk.END, f"Text: {tweet}")

def done():
    root.destroy()

def open_input_box(prompt, title="Input Box"):
    root = tk.Tk()
    root.withdraw()  # Hide the main tkinter window
    result = simpledialog.askstring(title, prompt)
    root.destroy()  # Close the hidden tkinter window
    return result

def change_font_and_size(docx_filename):
    try:
        # Open the document
        doc = Document(docx_filename)

        # Set the desired font name (e.g., Calibri) and font size (e.g., 11)
        font_name = 'Calibri'
        font_size = Pt(11)

        # Loop through paragraphs and change font and size
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = font_size

        # Save the modified document, overwriting the original
        doc.save(docx_filename)

    except Exception as e:
        print(f"An error occurred: {e}")

def format_tweets(tweets, urls, comments):
    formatted_tweets = []
    for tweet, url, comment in zip(tweets, urls, comments):
        # Extract date, author, and full text using regular expressions
        match = re.match(r'(.+?) @(.+?) · (\w+ \d+) (.+)', tweet)
        if match:
            author, date, full_text = match.group(1), match.group(3), match.group(4)
            # Extract the day and month from the date
            month, day = re.match(r'(\w+) (\d+)', date).groups()
            # Convert the month name to a month number
            formatted_month = month_to_number.get(month, month)
            # Format the date as "DD/MM/YYYY"
            formatted_date = f"{day}/{formatted_month}/2023"
            formatted_tweet = f"{formatted_date} {author} tweets \"{full_text}\""
            if comment:  # If a comment is provided, append it to the formatted tweet
                formatted_tweet += f", Comment: {comment}"
            if url:  # If a URL is provided, append it to the formatted tweet
                formatted_tweet += f", URL: {url}"
            formatted_tweets.append(formatted_tweet)
    print('Printing formatted data...')
    for formatted_tweet in formatted_tweets:
        print(formatted_tweet)
    return formatted_tweets

tweets = []
urls = []
comments = []  # Store comments

# Uses tkinter to create a GUI to collect tweets, comments, and URLs
root = tk.Tk()
root.title("Data Entry")
tweet_entry_label = tk.Label(root, text="Paste tweet content and user name")
tweet_entry_label.pack()
tweet_entry = tk.Entry(root, width=40)
tweet_entry.pack()
url_entry_label = tk.Label(root, text="Enter URL (optional)")
url_entry_label.pack()
url_entry = tk.Entry(root, width=40)
url_entry.pack()
comment_entry_label = tk.Label(root, text="Add comment (optional)")
comment_entry_label.pack()
comment_entry = tk.Entry(root, width=40)
comment_entry.pack()
add_tweet_button = tk.Button(root, text="Add Tweet", command=add_tweet)
add_tweet_button.pack()

# Place the "Double click to remove entry" label below the listbox
double_click_label = tk.Label(root, text="Double click to remove entry")
double_click_label.pack()

tweet_listbox = tk.Listbox(root, height=10, selectmode=tk.SINGLE)
tweet_listbox.pack()

done_button = tk.Button(root, text="Done", command=done)
done_button.pack()

# Bind the double-click event to the delete_selected_item function
tweet_listbox.bind("<Double-Button-1>", delete_selected_item)

root.mainloop()

print("Captured Data:")
print("Tweets:", tweets)
print("Comments:", comments)

# Pre-cleaning text for regex
for i in range(len(tweets)):
    tweets[i] = tweets[i].replace('\n', " ")  # Clean tweet text
print("Cleaned Data:")
print(tweets)

# Clean and format data
write_data = format_tweets(tweets, urls, comments)

# Write to document
doc = Document()
for tweet in write_data:
    doc.add_paragraph(tweet, style='List Bullet')

file_name = open_input_box("Subject name & Date", "Taji Mustafa Oct 8")
file_name = file_name + '.docx'
doc.save(file_name)
print('Formatting file...')
change_font_and_size(file_name)
print('Done!')

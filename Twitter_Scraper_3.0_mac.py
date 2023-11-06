import tkinter as tk
from tkinter import simpledialog
from docx import Document
from docx.shared import Pt
import os
import sys
from selenium import webdriver
from selenium.webdriver.firefox.service import Service
from selenium.webdriver.firefox.options import Options
import time
import re

month_to_number = {
    'Jan': '01', 'Feb': '02', 'Mar': '03', 'Apr': '04',
    'May': '05', 'Jun': '06', 'Jul': '07', 'Aug': '08',
    'Sep': '09', 'Oct': '10', 'Nov': '11', 'Dec': '12'
}

data_list = []

gecko_driver_path = '/usr/local/bin/geckodriver'


def archive_url(saved_url):
    # Set up the Firefox webdriver with the service
    service = Service(gecko_driver_path)
    driver = webdriver.Firefox(service=service)
    # Archives page
    driver.get('https://archive.ph/?run=1&url='+ saved_url)
    # Wait for 3 seconds for the page to load/to archive
    time.sleep(3)
    # Get the current URL
    current_url = driver.current_url
    # Close the browser
    driver.quit()
    return current_url

def replace_text_in_docx(docx_path):
    # Load the DOCX file
    doc = Document(docx_path)

    # Define a regular expression pattern to match the desired text
    pattern = r'@[A-Za-z0-9]+ · \d{1,2}h'

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Search for the pattern in the text
            matches = re.findall(pattern, run.text)
            for match in matches:
                # Replace the matched text with 'tweets'
                run.text = run.text.replace(match, 'tweets:')

    # Save the modified DOCX file
    doc.save(docx_path)


def get_script_directory():
    # Get the directory where the script is located
    script_directory = os.path.dirname(os.path.abspath(__file__))
    return script_directory

def save_document(doc, filename):
    # Get the script's directory and create the 'output' directory if it doesn't exist
    script_directory = get_script_directory()
    output_directory = os.path.join(script_directory, 'output')
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    # Create the full file path within the 'output' directory
    full_path = os.path.join(output_directory, filename)
    # Save the document in the 'output' directory
    doc.save(full_path)
    return full_path

def add_tweet():
    tweet = tweet_entry.get()
    url = url_entry.get()  # Get the URL input
    alternative_format = alternative_format_entry.get()  # Get the alternative format input
    if tweet:
        # Remove newlines from the tweet
        tweet = tweet.replace('\n', ' ')
        data_list.append((tweet, url, alternative_format))  # Store the data with input order
        tweet_entry.delete(0, tk.END)
        url_entry.delete(0, tk.END)  # Clear the URL input field
        alternative_format_entry.delete(0, tk.END)  # Clear the alternative format input field
        update_tweet_list()

def delete_selected_item(event):
    selected_index = tweet_listbox.curselection()  # Get the index of the selected item(s)
    if selected_index:
        for index in selected_index:
            index = int(index)
            del data_list[index]  # Delete the selected item from the data list
            tweet_listbox.delete(index)  # Delete the selected item from the Listbox
            update_tweet_list()  # Update the Listbox to reflect the changes

def update_tweet_list():
    tweet_listbox.delete(0, tk.END)
    for i, data in enumerate(data_list):
        tweet_listbox.insert(tk.END, f"Text: {data[0]}")

def done():
    root.destroy()

def open_input_box(prompt, title="Input Box"):
    root = tk.Tk()
    root.withdraw()  # Hide the main tkinter window
    result = simpledialog.askstring(title, prompt)
    root.destroy()  # Close the hidden tkinter window
    return result

def change_font_and_size(docx_filename):
    try:
        # Open the document
        doc = Document(docx_filename)

        # Set the desired font name (e.g., Calibri) and font size (e.g., 11)
        font_name = 'Calibri'
        font_size = Pt(11)

        # Loop through paragraphs and change font and size
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = font_size

        # Save the modified document, overwriting the original
        doc.save(docx_filename)

    except Exception as e:
        print(f"An error occurred: {e}")

def format_tweets(data_list):
    formatted_tweets = []

    for data in data_list:
        tweet, url, alternative_format = data
        # Check if the data matches the new format
        if re.match(r'(.+?) @(.+?) · (\w+ \d+) (.+)', tweet):
            # Extract date, author, and full text using regular expressions
            match = re.match(r'(.+?) @(.+?) · (\w+ \d+) (.+)', tweet)
            if match:
                author, date, full_text = match.group(1), match.group(3), match.group(4)
                # Extract the day and month from the date
                month, day = re.match(r'(\w+) (\d+)', date).groups()
                # Convert the month name to a month number
                formatted_month = month_to_number.get(month, month)
                # Format the date as "DD/MM/YYYY"
                formatted_date = f"{day}/{formatted_month}/2023"
                formatted_tweet = f"{formatted_date} {author} tweets \"{full_text}\""
                if alternative_format:  # If an alternative format is provided, append it to the formatted tweet
                    formatted_tweet += f", Comment: {alternative_format}"
                if url:  # If a URL is provided, append it to the formatted tweet
                    formatted_tweet += f", URL: {url}"
                    print("Archiving page....")
                    archived_url = archive_url(url)
                    formatted_tweet += f", Archive: {archived_url}"
                    print("Done!")
                formatted_tweets.append(formatted_tweet)
        else:
            different_format_data = tweet
            if alternative_format:  # If an alternative format is provided, append it to the different format data
                different_format_data += f", Comment: {alternative_format}"
            if url:  # If a URL is provided, append it to the different format data
                different_format_data += f", URL: {url}"
                print("Archiving page....")
                archived_url = archive_url(url)
                different_format_data += f", Archive: {archived_url}"
                print("Done!")
            formatted_tweets.append(different_format_data)

    return formatted_tweets

# Uses tkinter to create a GUI to collect tweets, alternative formats, and URLs
root = tk.Tk()
root.title("Data Entry")
tweet_entry_label = tk.Label(root, text="Paste tweet content and user name")
tweet_entry_label.pack()
tweet_entry = tk.Entry(root, width=40)
tweet_entry.pack()
url_entry_label = tk.Label(root, text="Enter URL (optional)")
url_entry_label.pack()
url_entry = tk.Entry(root, width=40)
url_entry.pack()
alternative_format_label = tk.Label(root, text="Comment (optional)")
alternative_format_label.pack()
alternative_format_entry = tk.Entry(root, width=40)
alternative_format_entry.pack()
add_tweet_button = tk.Button(root, text="Add Tweet", command=add_tweet)
add_tweet_button.pack()

# Place the "Double click to remove entry" label below the listbox
double_click_label = tk.Label(root, text="Double click to remove entry")
double_click_label.pack()

tweet_listbox = tk.Listbox(root, height=10, selectmode=tk.SINGLE)
tweet_listbox.pack()

done_button = tk.Button(root, text="Done", command=done)
done_button.pack()

# Bind the double-click event to the delete_selected_item function
tweet_listbox.bind("<Double-Button-1>", delete_selected_item)

root.mainloop()

print("Captured Data:")
print("Data List:", data_list)

# Clean and format data
write_data = format_tweets(data_list)

# Prompt the user for the subject name and date
subject_name_and_date = open_input_box("Subject name & Date", "Subject name & Date")
file_name = subject_name_and_date + '.docx'

# Create the Word document
doc = Document()
for data in write_data:
    doc.add_paragraph(data, style='List Bullet')

# After creating the Word document, use the save_document function to save it in the script's directory
saved_file_path = save_document(doc, file_name)
print(f"Saved document to: {saved_file_path}")

# Change the font and size of the saved document
print('Formatting file...')
change_font_and_size(saved_file_path)
replace_text_in_docx(saved_file_path)
print('Done!')

input("Press Enter to exit...")
